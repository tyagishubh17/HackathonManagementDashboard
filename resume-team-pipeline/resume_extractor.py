import os
import pdfplumber
import docx

try:

    import pytesseract
    from pdf2image import convert_from_path

    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

def _ocr_pdf_page(filepath: str, page_number: int) -> str:
    if not OCR_AVAILABLE:
        return ""
    try:
        images = convert_from_path(
            filepath, first_page=page_number + 1, last_page=page_number + 1, dpi=300
        )
        if not images:
            return ""
        return pytesseract.image_to_string(images[0])
    except Exception as e:
        print(f"  [WARN] OCR failed on page {page_number + 1} of {filepath}: {e}")
        return ""

def extract_text_from_pdf(filepath: str) -> str:
    text_chunks = []
    ocr_used = False
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                text_chunks.append(page_text)
            else:
                ocr_text = _ocr_pdf_page(filepath, i)
                if ocr_text.strip():
                    text_chunks.append(ocr_text)
                    ocr_used = True
    if ocr_used:
        print(f"  [INFO] Used OCR fallback for one or more pages in {filepath}")
    return "\n".join(text_chunks)

def extract_text_from_docx(filepath: str) -> str:
    document = docx.Document(filepath)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)

def extract_resume_text(filepath: str) -> str:
    if not os.path.exists(filepath):
        print(f"  [WARN] File not found: {filepath}")
        return ""
    ext = os.path.splitext(filepath)[1].lower()
    try:
        if ext == ".pdf":
            return extract_text_from_pdf(filepath)
        elif ext == ".docx":
            return extract_text_from_docx(filepath)
        else:
            print(f"  [WARN] Unsupported file type '{ext}' for {filepath}")
            return ""
    except Exception as e:
        print(f"  [WARN] Failed to parse {filepath}: {e}")
        return ""
